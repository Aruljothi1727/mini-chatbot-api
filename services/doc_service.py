import io
import re
import pdfplumber
import docx
from langchain_core.documents import Document
from typing import List, Tuple

class DocumentChunker:
    def __init__(self): 
        self.item_code_pattern = re.compile(r'^[A-Z]\d{4}[A-Z]?:', re.MULTILINE)
        self.section_pattern = re.compile(r'^SECTION [A-Z]{1,2}:', re.MULTILINE)
        self.chapter_pattern = re.compile(r'^CHAPTER \d+:', re.MULTILINE)
        self.subsection_pattern = re.compile(r'^\d+\.\d+(\.\d+)*\s+', re.MULTILINE)
        print("Initial DocumentChunker with regex patterns.\n")

    def extract_text_with_structure(self, file_bytes: bytes, filename: str) -> str:
        content = ""
        print(f"--- Extracting text from {filename} ---")

        if filename.endswith(".pdf"):
            pdf_file = io.BytesIO(file_bytes)
            with pdfplumber.open(pdf_file) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    content += f"\n--- PAGE {page_num} ---\n{text}\n"
                    print(f"Extracted page {page_num} text ({len(text)} chars).")

        elif filename.endswith(".docx"):
            doc_file = io.BytesIO(file_bytes)
            document = docx.Document(doc_file)
            for i, para in enumerate(document.paragraphs, 1):
                if para.text.strip():
                    content += para.text.strip() + "\n"
                    print(f"Extracted paragraph {i} ({len(para.text.strip())} chars).")

        print(f"Total extracted text length: {len(content)} chars.\n")
        return content

    def find_all_structural_markers(self, content: str) -> List[Tuple[int, str, str]]:
        markers = []
        print("--- Finding structural markers ---")

        for pattern, mtype in [(self.chapter_pattern, 'CHAPTER'),
                               (self.section_pattern, 'SECTION'),
                               (self.item_code_pattern, 'ITEM'),
                               (self.subsection_pattern, 'SUBSECTION')]:
            for match in pattern.finditer(content):
                line = content[match.start():content.find('\n', match.start())].strip()
                markers.append((match.start(), mtype, line))
                print(f"Found {mtype}: '{line}' at position {match.start()}")

        markers.sort(key=lambda x: x[0])
        print(f"Total markers found: {len(markers)}\n")
        return markers

    def create_hierarchical_chunks(self, content: str) -> List[Document]:
        markers = self.find_all_structural_markers(content)
        chunks = []
        current_context = {'chapter': '', 'section': '', 'subsection': ''}

        print("--- Creating hierarchical chunks ---")
        for i, (pos, marker_type, marker_text) in enumerate(markers):
            if marker_type == 'CHAPTER':
                current_context['chapter'] = marker_text
                current_context['section'] = ''
                current_context['subsection'] = ''
            elif marker_type == 'SECTION':
                current_context['section'] = marker_text
                current_context['subsection'] = ''
            elif marker_type == 'SUBSECTION':
                current_context['subsection'] = marker_text

            start_pos = pos
            end_pos = markers[i + 1][0] if i + 1 < len(markers) else len(content)
            chunk_content = content[start_pos:end_pos].strip()
            if len(chunk_content) < 50:
                print(f"Skipped short chunk at position {start_pos} ({len(chunk_content)} chars).")
                continue

            # ITEM chunks
            if marker_type == 'ITEM':
                item_code_match = re.match(r'^([A-Z]\d{4}[A-Z]?):', marker_text)
                item_code = item_code_match.group(1) if item_code_match else 'UNKNOWN'
                context_header = f"Context: {current_context['chapter']} > {current_context['section']} > {current_context['subsection']}\n"
                context_header += f"Item Code: {item_code}\n" + "-" * 50 + "\n"
                chunk_content = context_header + chunk_content
                chunks.append(Document(
                    page_content=chunk_content,
                    metadata={
                        'chunk_type': 'item',
                        'item_code': item_code,
                        'chapter': current_context['chapter'],
                        'section': current_context['section'],
                        'subsection': current_context['subsection'],
                        'chunk_id': len(chunks)
                    }
                ))
                print(f"Created ITEM chunk: {item_code} | Context: {current_context} | Length: {len(chunk_content)} chars")

            # SECTION / CHAPTER chunks
            elif marker_type in ['SECTION', 'CHAPTER']:
                preview_content = chunk_content[:2000] + "..." if len(chunk_content) > 2000 else chunk_content
                chunks.append(Document(
                    page_content=preview_content,
                    metadata={
                        'chunk_type': marker_type.lower(),
                        'title': marker_text,
                        'chapter': current_context['chapter'],
                        'section': current_context['section'] if marker_type != 'SECTION' else marker_text,
                        'chunk_id': len(chunks)
                    }
                ))
                print(f"Created {marker_type} chunk: '{marker_text}' | Context: {current_context} | Length: {len(preview_content)} chars")

        print(f"Total hierarchical chunks created: {len(chunks)}\n")
        return chunks

    def create_overlapping_windows(self, chunks: List[Document], window_size: int = 1500, overlap: int = 200) -> List[Document]:
        windowed_chunks = []
        print("--- Creating overlapping windows ---")

        for chunk in chunks:
            content = chunk.page_content
            if len(content) <= window_size:
                windowed_chunks.append(chunk)
                continue

            start = 0
            chunk_num = 0
            while start < len(content):
                end = min(start + window_size, len(content))
                window_content = content[start:end]

                if end < len(content):
                    last_period = window_content.rfind('.')
                    last_newline = window_content.rfind('\n')
                    break_point = max(last_period, last_newline)
                    if break_point > start + window_size * 0.7:
                        end = break_point + 1
                        window_content = content[start:end]

                new_metadata = chunk.metadata.copy()
                new_metadata['window_num'] = chunk_num
                new_metadata['chunk_id'] = f"{chunk.metadata.get('chunk_id', 0)}_{chunk_num}"
                windowed_chunks.append(Document(page_content=window_content.strip(), metadata=new_metadata))

                print(f"Chunk {chunk.metadata.get('chunk_id', 'unknown')} window {chunk_num}: length {len(window_content.strip())} chars")

                start = end - overlap
                chunk_num += 1
                if start >= len(content) - overlap:
                    break

        print(f"Total windowed chunks created: {len(windowed_chunks)}\n")
        return windowed_chunks

    def process_file(self, file_bytes: bytes, filename: str) -> List[Document]:
        print(f"=== Processing {filename} ===\n")
        content = self.extract_text_with_structure(file_bytes, filename)
        chunks = self.create_hierarchical_chunks(content)
        windowed_chunks = self.create_overlapping_windows(chunks)
        filtered_chunks = [c for c in windowed_chunks if len(c.page_content.strip()) >= 100]

        print(f"=== Finished processing ===")
        print(f"Initial chunks: {len(chunks)} | Windowed chunks: {len(windowed_chunks)} | Filtered chunks: {len(filtered_chunks)}\n")

        # Show first few chunks for verification
        for i, chunk in enumerate(filtered_chunks[:3]):
            print(f"Sample Chunk {i+1}:")
            print(f"Type: {chunk.metadata.get('chunk_type', 'unknown')}")
            print(f"Item Code: {chunk.metadata.get('item_code', 'N/A')}")
            print(f"Context: Chapter='{chunk.metadata.get('chapter')}', Section='{chunk.metadata.get('section')}', Subsection='{chunk.metadata.get('subsection', '')}'")
            print(f"Length: {len(chunk.page_content)} chars")
            print(f"Preview:\n{chunk.page_content[:300]}...")
            print("-" * 50)

        return filtered_chunks


# Usage function
def process_document(file_bytes: bytes, filename: str) -> List[Document]:
    chunker = DocumentChunker()
    return chunker.process_file(file_bytes, filename)

